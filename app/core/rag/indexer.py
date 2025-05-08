"""
Indexer Module for CasaLingua
Provides functionality for indexing, chunking, and preparing documents for the RAG system
"""

import os
import json
import hashlib
import re
import logging
import uuid
import tempfile
from typing import List, Dict, Union, Tuple, Optional, Any, Callable, Iterator, BinaryIO
from datetime import datetime
from docx import Document
import pandas as pd
from tqdm import tqdm

# Tokenizer pipeline import
from app.core.pipeline.tokenizer import TokenizerPipeline
# ModelRegistry import for dynamic tokenizer loading
from app.services.models.loader import ModelRegistry
# Session manager for accessing user documents
from app.services.storage.session_manager import SessionManager
# Utils import
from app.utils.logging import get_logger

# Get logger
logger = get_logger(__name__)

class DocumentProcessor:
    """Base class for document processors"""
    
    def __init__(self):
        """Initialize the document processor"""
        pass
    
    def process(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Process a document into chunks suitable for indexing
        
        Args:
            file_path (str): Path to the document file
            
        Returns:
            List[Dict[str, Any]]: List of document chunks with metadata
        """
        raise NotImplementedError("Subclasses must implement process()")
    
    @staticmethod
    def detect_language(text: str) -> str:
        """
        Simple keyword-based heuristic for language detection.
        Replace with langdetect or spaCy for production quality.

        Args:
            text (str): Text to detect language for

        Returns:
            str: Detected language code (e.g., 'en', 'es', 'fr')
        """
        text = text.lower()
        language_markers = {
            'en': ['the', 'and', 'of', 'to', 'in', 'is', 'you', 'that', 'it', 'he'],
            'es': ['el', 'la', 'de', 'que', 'y', 'a', 'en', 'un', 'ser', 'se'],
            'fr': ['le', 'la', 'de', 'et', 'est', 'en', 'que', 'un', 'une', 'du'],
            'de': ['der', 'die', 'das', 'und', 'ist', 'in', 'zu', 'den', 'mit', 'nicht'],
            'it': ['il', 'la', 'di', 'e', 'che', 'un', 'a', 'per', 'in', 'sono']
        }
        scores = {lang: 0 for lang in language_markers}
        for lang, markers in language_markers.items():
            for marker in markers:
                pattern = r'\b' + re.escape(marker) + r'\b'
                matches = re.findall(pattern, text)
                scores[lang] += len(matches)
        max_score = max(scores.values())
        if max_score == 0:
            return 'unknown'
        return max(scores.items(), key=lambda x: x[1])[0]


class TextProcessor(DocumentProcessor):
    """Processor for plain text files"""
    
    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50, tokenizer: Optional[TokenizerPipeline] = None):
        """
        Initialize the text processor
        
        Args:
            chunk_size (int): Maximum size of text chunks in characters
            chunk_overlap (int): Overlap between consecutive chunks in characters
            tokenizer (TokenizerPipeline, optional): Shared tokenizer pipeline
        """
        super().__init__()
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.tokenizer = tokenizer
    
    def process(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Process a text file into chunks
        
        Args:
            file_path (str): Path to the text file
            
        Returns:
            List[Dict[str, Any]]: List of text chunks with metadata
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
                
            return self._chunk_text(text, os.path.basename(file_path))
                
        except Exception as e:
            logger.error(f"Error processing text file {file_path}: {e}")
            return []
    
    def _chunk_text(self, text: str, source: str) -> List[Dict[str, Any]]:
        """
        Split text into overlapping chunks
        
        Args:
            text (str): Text to split
            source (str): Source identifier
            
        Returns:
            List[Dict[str, Any]]: List of text chunks with metadata
        """
        chunks = []
        
        # Simple chunking by character count
        # For better results, consider chunking by sentences or paragraphs
        start = 0
        while start < len(text):
            end = min(start + self.chunk_size, len(text))
            
            # If we're not at the end, try to break at a sentence boundary
            if end < len(text):
                # Look for sentence boundaries (., !, ?)
                last_period = max(
                    text.rfind('. ', start, end),
                    text.rfind('! ', start, end),
                    text.rfind('? ', start, end)
                )
                
                if last_period != -1:
                    end = last_period + 1  # Include the period
            
            chunk_text = text[start:end].strip()
            
            if chunk_text:  # Only add non-empty chunks
                # Detect language
                language = self.detect_language(chunk_text)
                
                # Create unique ID for the chunk
                chunk_id = hashlib.md5(chunk_text.encode('utf-8')).hexdigest()
                chunk = {
                    "id": chunk_id,
                    "text": chunk_text,
                    "source": source,
                    "metadata": {
                        "language": language,
                        "start_char": start,
                        "end_char": end,
                        "character_count": len(chunk_text)
                    }
                }
                if self.tokenizer:
                    chunk_tokens = self.tokenizer.encode(chunk_text)
                    chunk["metadata"]["tokens"] = chunk_tokens
                chunks.append(chunk)
            
            # Move to next chunk with overlap
            start = end - self.chunk_overlap if end < len(text) else end
        
        return chunks


class DocxProcessor(DocumentProcessor):
    """Processor for Microsoft Word documents"""
    
    def __init__(self, max_paragraphs_per_chunk: int = 5, tokenizer: Optional[TokenizerPipeline] = None):
        """
        Initialize the Word document processor
        
        Args:
            max_paragraphs_per_chunk (int): Maximum number of paragraphs per chunk
            tokenizer (TokenizerPipeline, optional): Shared tokenizer pipeline
        """
        super().__init__()
        self.max_paragraphs_per_chunk = max_paragraphs_per_chunk
        self.tokenizer = tokenizer
    
    def process(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Process a Word document into chunks
        
        Args:
            file_path (str): Path to the Word document
            
        Returns:
            List[Dict[str, Any]]: List of document chunks with metadata
        """
        try:
            doc = Document(file_path)
            
            chunks = []
            current_paragraphs = []
            current_text = ""
            
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():  # Skip empty paragraphs
                    current_paragraphs.append(paragraph.text)
                    current_text += paragraph.text + "\n"
                    
                    # Create chunk when we reach max paragraphs
                    if len(current_paragraphs) >= self.max_paragraphs_per_chunk:
                        chunk_text = "\n".join(current_paragraphs)
                        
                        # Detect language
                        language = self.detect_language(chunk_text)
                        
                        # Create unique ID for the chunk
                        chunk_id = hashlib.md5(chunk_text.encode('utf-8')).hexdigest()
                        chunk = {
                            "id": chunk_id,
                            "text": chunk_text,
                            "source": os.path.basename(file_path),
                            "metadata": {
                                "language": language,
                                "paragraph_index_range": f"{i-len(current_paragraphs)+1}-{i}",
                                "paragraph_count": len(current_paragraphs),
                                "character_count": len(chunk_text)
                            }
                        }
                        if self.tokenizer:
                            chunk_tokens = self.tokenizer.encode(chunk_text)
                            chunk["metadata"]["tokens"] = chunk_tokens
                        chunks.append(chunk)
                        
                        current_paragraphs = []
                        current_text = ""
            
            # Add remaining paragraphs as a chunk
            if current_paragraphs:
                chunk_text = "\n".join(current_paragraphs)
                
                # Detect language
                language = self.detect_language(chunk_text)
                
                # Create unique ID for the chunk
                chunk_id = hashlib.md5(chunk_text.encode('utf-8')).hexdigest()
                chunk = {
                    "id": chunk_id,
                    "text": chunk_text,
                    "source": os.path.basename(file_path),
                    "metadata": {
                        "language": language,
                        "paragraph_index_range": f"{len(doc.paragraphs)-len(current_paragraphs)+1}-{len(doc.paragraphs)}",
                        "paragraph_count": len(current_paragraphs),
                        "character_count": len(chunk_text)
                    }
                }
                if self.tokenizer:
                    chunk_tokens = self.tokenizer.encode(chunk_text)
                    chunk["metadata"]["tokens"] = chunk_tokens
                chunks.append(chunk)
            
            return chunks
                
        except Exception as e:
            logger.error(f"Error processing Word document {file_path}: {e}")
            return []


class CSVProcessor(DocumentProcessor):
    """Processor for CSV files"""
    
    def __init__(self, text_columns: Optional[List[str]] = None, tokenizer: Optional[TokenizerPipeline] = None):
        """
        Initialize the CSV processor
        
        Args:
            text_columns (List[str], optional): List of column names to include
            tokenizer (TokenizerPipeline, optional): Shared tokenizer pipeline
        """
        super().__init__()
        self.text_columns = text_columns
        self.tokenizer = tokenizer
    
    def process(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Process a CSV file into document chunks
        
        Args:
            file_path (str): Path to the CSV file
            
        Returns:
            List[Dict[str, Any]]: List of document chunks with metadata
        """
        try:
            # Read CSV file
            df = pd.read_csv(file_path)
            
            # If no specific text columns provided, use all columns
            text_columns = self.text_columns or df.columns.tolist()
            
            chunks = []
            
            # Process each row
            for i, row in df.iterrows():
                # Build text from specified columns
                row_text = " ".join(str(row[col]) for col in text_columns if col in row)
                
                if row_text.strip():
                    # Detect language
                    language = self.detect_language(row_text)
                    
                    # Create unique ID for the chunk
                    chunk_id = hashlib.md5(row_text.encode('utf-8')).hexdigest()
                    
                    # Standardized metadata keys
                    metadata = {
                        "row_index": i,
                        "language": language,
                        "character_count": len(row_text)
                    }
                    # Add any additional columns as metadata
                    for col in df.columns:
                        if col not in text_columns and pd.notna(row[col]):
                            metadata[col] = row[col]
                    if self.tokenizer:
                        chunk_tokens = self.tokenizer.encode(row_text)
                        metadata["tokens"] = chunk_tokens
                    chunks.append({
                        "id": chunk_id,
                        "text": row_text,
                        "source": os.path.basename(file_path),
                        "metadata": metadata
                    })
            
            return chunks
                
        except Exception as e:
            logger.error(f"Error processing CSV file {file_path}: {e}")
            return []


class Indexer:
    """Main indexer class for CasaLingua"""

    def __init__(self,
                output_dir: str = "./index",
                chunk_size: int = 500,
                chunk_overlap: int = 50,
                rag_expert=None,
                processor=None):
        """
        Initialize the indexer

        Args:
            output_dir (str): Directory to save indexed documents
            chunk_size (int): Default chunk size for text processors
            chunk_overlap (int): Default chunk overlap for text processors
            rag_expert: Optional RAG Expert instance to update with new documents
            processor: Optional unified processor for document handling
        """
        self.output_dir = output_dir
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.rag_expert = rag_expert
        self.processor = processor
        
        # Initialize session manager
        self.session_manager = SessionManager()

        # Load tokenizer dynamically from registry for RAG retriever task
        registry = ModelRegistry()
        _, tokenizer_name = registry.get_model_and_tokenizer("rag_retriever")
        tokenizer = TokenizerPipeline(model_name=tokenizer_name, task_type="rag_retrieval")

        # Initialize document processors with dynamically loaded tokenizer
        self.processors = {
            ".txt": TextProcessor(chunk_size, chunk_overlap, tokenizer=tokenizer),
            ".docx": DocxProcessor(tokenizer=tokenizer),
            ".csv": CSVProcessor(tokenizer=tokenizer)
        }

        # Create output directory if it doesn't exist
        os.makedirs(output_dir, exist_ok=True)
    
    def register_processor(self, extension: str, processor: DocumentProcessor) -> None:
        """
        Register a custom document processor
        
        Args:
            extension (str): File extension to associate with processor
            processor (DocumentProcessor): Processor instance
        """
        self.processors[extension.lower()] = processor
    
    def index_file(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Index a single file
        
        Args:
            file_path (str): Path to the file
            
        Returns:
            List[Dict[str, Any]]: List of indexed document chunks
        """
        if not os.path.isfile(file_path):
            logger.warning(f"File does not exist: {file_path}")
            return []
            
        # Get file extension
        ext = os.path.splitext(file_path)[1].lower()
        
        # Check if we have a processor for this extension
        if ext not in self.processors:
            logger.warning(f"No processor registered for extension {ext}")
            return []
            
        # Process the file
        logger.info(f"Indexing file: {file_path}")
        processor = self.processors[ext]
        chunks = processor.process(file_path)
        
        # Add indexing metadata
        for chunk in chunks:
            if "metadata" not in chunk:
                chunk["metadata"] = {}
                
            chunk["metadata"].update({
                "indexed_at": datetime.now().isoformat(),
                "file_path": file_path,
                "file_name": os.path.basename(file_path),
                "file_extension": ext
            })
        
        return chunks
    
    def index_directory(self, 
                      directory: str, 
                      recursive: bool = True,
                      file_extensions: Optional[List[str]] = None) -> List[Dict[str, Any]]:
        """
        Index all files in a directory
        
        Args:
            directory (str): Directory path
            recursive (bool): Whether to process subdirectories
            file_extensions (List[str], optional): Only process these extensions
            
        Returns:
            List[Dict[str, Any]]: List of all indexed document chunks
        """
        if not os.path.isdir(directory):
            logger.warning(f"Directory does not exist: {directory}")
            return []
            
        all_chunks = []
        
        # Get list of files
        files = []
        if recursive:
            for root, _, filenames in os.walk(directory):
                for filename in filenames:
                    files.append(os.path.join(root, filename))
        else:
            files = [os.path.join(directory, f) for f in os.listdir(directory) 
                    if os.path.isfile(os.path.join(directory, f))]
        
        # Filter by extension if needed
        if file_extensions:
            filtered_extensions = [ext.lower() if ext.startswith('.') else f".{ext.lower()}" 
                                for ext in file_extensions]
            files = [f for f in files if os.path.splitext(f)[1].lower() in filtered_extensions]
        
        # Process each file
        for file_path in tqdm(files, desc="📁 Indexing files"):
            chunks = self.index_file(file_path)
            all_chunks.extend(chunks)
        
        return all_chunks
    
    def save_index(self, chunks: List[Dict[str, Any]], output_file: Optional[str] = None) -> str:
        """
        Save indexed chunks to a file
        
        Args:
            chunks (List[Dict[str, Any]]): List of document chunks
            output_file (str, optional): Output file path
            
        Returns:
            str: Path to the saved file
        """
        if not chunks:
            logger.warning("No chunks to save")
            return ""
            
        if not output_file:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_file = os.path.join(self.output_dir, f"index_{timestamp}.json")
        
        try:
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(chunks, f, ensure_ascii=False, indent=2)
            logger.info(f"[Indexer] ✅ Saved {len(chunks)} chunks to '{output_file}'")
            
            # Update RAG expert's knowledge base if available
            if self.rag_expert:
                self.rag_expert.knowledge_base.extend(chunks)
                # Trigger index rebuild
                if hasattr(self.rag_expert, "_build_index"):
                    import asyncio
                    asyncio.create_task(self.rag_expert._build_index())
                
            return output_file
        except Exception as e:
            logger.error(f"Error saving index: {e}")
            return ""
    
    def index_and_save(self, 
                     path: str, 
                     output_file: Optional[str] = None,
                     recursive: bool = True,
                     file_extensions: Optional[List[str]] = None) -> Tuple[str, int]:
        """
        Index a file or directory and save the results
        
        Args:
            path (str): Path to file or directory
            output_file (str, optional): Output file path
            recursive (bool): Whether to process subdirectories
            file_extensions (List[str], optional): Only process these extensions
            
        Returns:
            Tuple[str, int]: Output file path and number of chunks indexed
        """
        chunks = []
        
        if os.path.isfile(path):
            chunks = self.index_file(path)
        elif os.path.isdir(path):
            chunks = self.index_directory(path, recursive, file_extensions)
        else:
            logger.warning(f"Path does not exist: {path}")
            return "", 0
        
        output_path = self.save_index(chunks, output_file)
        return output_path, len(chunks)
    
    async def index_document_content(self,
                                    document_content: bytes,
                                    document_type: str,
                                    filename: str,
                                    metadata: Optional[Dict[str, Any]] = None) -> List[Dict[str, Any]]:
        """
        Index document content directly from bytes
        
        Args:
            document_content: The document content as bytes
            document_type: MIME type of the document
            filename: Original filename
            metadata: Additional metadata
            
        Returns:
            List of indexed document chunks
        """
        metadata = metadata or {}
        document_id = metadata.get("document_id", str(uuid.uuid4()))
        chunks = []
        
        try:
            # Extract text based on document type
            extracted_text = ""
            
            if self.processor:
                # Use the unified processor if available
                extraction_result = await self.processor.extract_document_text(
                    document_content=document_content,
                    document_type=document_type,
                    options={"ocr_enabled": True},
                    filename=filename,
                    request_id=str(uuid.uuid4())
                )
                extracted_text = extraction_result.get("text", "")
                
                # Add processor extraction metadata
                for key, value in extraction_result.get("metadata", {}).items():
                    metadata[f"extraction_{key}"] = value
            else:
                # Fallback to basic extraction
                # Save to temporary file for processors that need file paths
                with tempfile.NamedTemporaryFile(delete=False) as temp_file:
                    temp_file.write(document_content)
                    temp_path = temp_file.name
                
                try:
                    # Get file extension
                    ext = os.path.splitext(filename)[1].lower()
                    if ext in self.processors:
                        # Use registered processor
                        chunks = self.processors[ext].process(temp_path)
                        # Extract text from chunks for further processing
                        extracted_text = "\n\n".join([chunk.get("text", "") for chunk in chunks])
                    else:
                        # Try basic text decoding
                        try:
                            extracted_text = document_content.decode('utf-8')
                        except UnicodeDecodeError:
                            logger.warning(f"Unable to decode document: {filename}")
                finally:
                    # Clean up temp file
                    try:
                        os.unlink(temp_path)
                    except:
                        pass
            
            # If we got text but no chunks, create chunks
            if extracted_text and not chunks:
                text_processor = TextProcessor(self.chunk_size, self.chunk_overlap)
                chunks = text_processor._chunk_text(extracted_text, filename)
            
            # Add metadata to chunks
            for chunk in chunks:
                if "metadata" not in chunk:
                    chunk["metadata"] = {}
                    
                # Add document and indexing metadata
                chunk["metadata"].update({
                    "indexed_at": datetime.now().isoformat(),
                    "document_id": document_id,
                    "file_name": filename,
                    "mime_type": document_type
                })
                
                # Add any additional metadata
                for key, value in metadata.items():
                    chunk["metadata"][key] = value
            
            logger.info(f"Indexed document {filename} into {len(chunks)} chunks")
            return chunks
            
        except Exception as e:
            logger.error(f"Error indexing document content: {e}", exc_info=True)
            return []
    
    async def index_session_document(self,
                                    session_id: str,
                                    document_id: str,
                                    output_file: Optional[str] = None) -> Tuple[str, int]:
        """
        Index a document from a user session
        
        Args:
            session_id: Session identifier
            document_id: Document identifier
            output_file: Optional output file path
            
        Returns:
            Tuple with output file path and number of chunks
        """
        try:
            # Get document from session
            document_data = await self.session_manager.get_document(session_id, document_id)
            if not document_data:
                logger.warning(f"Document not found: {document_id} in session {session_id}")
                return "", 0
            
            # Extract metadata and content
            content = document_data.get("content", b"")
            metadata = document_data.get("metadata", {})
            filename = metadata.get("filename", f"document_{document_id}")
            document_type = metadata.get("content_type", "application/octet-stream")
            
            # Index document content
            chunks = await self.index_document_content(
                document_content=content,
                document_type=document_type,
                filename=filename,
                metadata={
                    "document_id": document_id,
                    "session_id": session_id,
                    **metadata
                }
            )
            
            # Save chunks to file
            output_path = self.save_index(chunks, output_file)
            return output_path, len(chunks)
            
        except Exception as e:
            logger.error(f"Error indexing session document: {e}", exc_info=True)
            return "", 0
    
    async def index_session_documents(self,
                                     session_id: str,
                                     document_ids: Optional[List[str]] = None,
                                     output_file: Optional[str] = None) -> Tuple[str, Dict[str, Any]]:
        """
        Index multiple documents from a user session
        
        Args:
            session_id: Session identifier
            document_ids: Optional list of document IDs (if None, index all)
            output_file: Optional output file path
            
        Returns:
            Tuple with output file path and results dictionary
        """
        all_chunks = []
        results = {
            "indexed_documents": 0,
            "total_chunks": 0,
            "documents": {}
        }
        
        try:
            # Get all documents from session
            if document_ids:
                documents = []
                for doc_id in document_ids:
                    doc = await self.session_manager.get_document(session_id, doc_id)
                    if doc:
                        documents.append(doc)
            else:
                # Get metadata for all documents in session
                documents_meta = await self.session_manager.get_all_documents(session_id)
                documents = []
                for doc_meta in documents_meta:
                    doc_id = doc_meta.get("document_id")
                    if doc_id:
                        doc = await self.session_manager.get_document(session_id, doc_id)
                        if doc:
                            documents.append(doc)
            
            # Process each document
            for document in documents:
                doc_id = document.get("document_id")
                try:
                    # Extract metadata and content
                    content = document.get("content", b"")
                    metadata = document.get("metadata", {})
                    filename = metadata.get("filename", f"document_{doc_id}")
                    document_type = metadata.get("content_type", "application/octet-stream")
                    
                    # Index document content
                    chunks = await self.index_document_content(
                        document_content=content,
                        document_type=document_type,
                        filename=filename,
                        metadata={
                            "document_id": doc_id,
                            "session_id": session_id,
                            **metadata
                        }
                    )
                    
                    # Add chunks to collection
                    all_chunks.extend(chunks)
                    
                    # Update results
                    results["indexed_documents"] += 1
                    results["total_chunks"] += len(chunks)
                    results["documents"][doc_id] = {
                        "chunks": len(chunks),
                        "filename": filename,
                        "document_type": document_type
                    }
                    
                except Exception as e:
                    logger.error(f"Error indexing document {doc_id}: {e}", exc_info=True)
                    results["documents"][doc_id] = {
                        "error": str(e),
                        "chunks": 0
                    }
            
            # Save all chunks to file
            if all_chunks:
                output_path = self.save_index(all_chunks, output_file)
                results["output_file"] = output_path
                results["status"] = "success"
            else:
                results["status"] = "no_chunks"
                results["output_file"] = ""
                output_path = ""
            
            return output_path, results
            
        except Exception as e:
            logger.error(f"Error indexing session documents: {e}", exc_info=True)
            return "", {
                "status": "error",
                "error": str(e),
                "indexed_documents": 0,
                "total_chunks": 0
            }


# Example usage
if __name__ == "__main__":
    # Create indexer with default settings
    indexer = Indexer(output_dir="./indexed_data")
    
    # Example 1: Index a single text file
    #file_chunks = indexer.index_file("path/to/language_lesson.txt")
    #indexer.save_index(file_chunks, "language_lesson_index.json")
    
    # Example 2: Index a directory
    #output_file, chunk_count = indexer.index_and_save(
    #    "path/to/language_materials",
    #    recursive=True,
    #    file_extensions=[".txt", ".docx", ".csv"]
    #)
    #print(f"Indexed {chunk_count} chunks, saved to {output_file}")
    
    # Example 3: Custom processor
    # class PDFProcessor(DocumentProcessor):
    #     def process(self, file_path: str) -> List[Dict[str, Any]]:
    #         # PDF processing logic here
    #         pass
    # 
    # indexer.register_processor(".pdf", PDFProcessor())
    
    print("Indexer module initialized. Use indexer.index_and_save() to begin indexing.")